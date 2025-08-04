import re
import logging
from typing import Dict, List, Optional, Tuple
from docx import Document
from dataclasses import dataclass
import os

from src.parsers.itmo_parser import Course, Program

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocxParser:
    def __init__(self):
        """Инициализация парсера Word документов"""
        self.current_program = None
        
    def parse_docx_file(self, file_path: str) -> Optional[Program]:
        """Парсинг Word документа с программой"""
        if not os.path.exists(file_path):
            logger.error(f"Файл не найден: {file_path}")
            return None
            
        try:
            doc = Document(file_path)
            logger.info(f"Парсинг файла: {file_path}")
            
            # Извлекаем весь текст из документа
            full_text = ""
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    full_text += text + "\n"
                    paragraphs.append(text)
            
            # Также читаем таблицы если есть
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            # Определяем программу по содержимому
            program_info = self._extract_program_info(full_text, paragraphs, tables_data, file_path)
            
            return program_info
            
        except Exception as e:
            logger.error(f"Ошибка при парсинге файла {file_path}: {e}")
            return None
    
    def _extract_program_info(self, full_text: str, paragraphs: List[str], 
                            tables_data: List, file_path: str) -> Program:
        """Извлечение информации о программе из содержимого документа"""
        
        # Определяем название программы
        program_name = self._extract_program_name(full_text, file_path)
        
        # Извлекаем описание
        description = self._extract_description(paragraphs)
        
        # Извлекаем продолжительность
        duration = self._extract_duration(full_text)
        
        # Извлекаем требования к поступлению
        admission_requirements = self._extract_admission_requirements(full_text, paragraphs)
        
        # Извлекаем карьерные перспективы
        career_prospects = self._extract_career_prospects(full_text, paragraphs)
        
        # Извлекаем курсы из таблиц и текста
        courses = self._extract_courses_from_document(full_text, paragraphs, tables_data, program_name)
        
        return Program(
            name=program_name,
            description=description,
            duration=duration,
            courses=courses,
            admission_requirements=admission_requirements,
            career_prospects=career_prospects
        )
    
    def _extract_program_name(self, full_text: str, file_path: str) -> str:
        """Извлечение названия программы"""
        
        # Ключевые слова для поиска названия программы
        ai_keywords = [
            'искусственный интеллект', 'artificial intelligence', 
            'машинное обучение', 'machine learning'
        ]
        
        ai_product_keywords = [
            'ai-продукт', 'ai продукт', 'ai product', 
            'продукт', 'разработка продуктов'
        ]
        
        text_lower = full_text.lower()
        
        # Определяем по ключевым словам
        if any(keyword in text_lower for keyword in ai_product_keywords):
            return "AI-продукты"
        elif any(keyword in text_lower for keyword in ai_keywords):
            return "Искусственный интеллект"
        
        # Определяем по номеру файла
        if '10033' in file_path:
            return "Искусственный интеллект"
        elif '10130' in file_path:
            return "AI-продукты"
        
        return "Неизвестная программа"
    
    def _extract_description(self, paragraphs: List[str]) -> str:
        """Извлечение описания программы"""
        
        # Ищем первый содержательный параграф
        for paragraph in paragraphs:
            if len(paragraph) > 100 and not self._is_header(paragraph):
                return paragraph
        
        # Если не найден отдельный описательный параграф, составляем из первых параграфов
        description_parts = []
        for paragraph in paragraphs[:5]:
            if len(paragraph) > 50 and not self._is_header(paragraph):
                description_parts.append(paragraph)
                if len(' '.join(description_parts)) > 200:
                    break
        
        return ' '.join(description_parts) if description_parts else "Описание программы"
    
    def _is_header(self, text: str) -> bool:
        """Проверка, является ли текст заголовком"""
        return (len(text) < 50 or 
                text.isupper() or 
                text.startswith(('№', 'Код', 'Название', 'Форма', 'Срок')))
    
    def _extract_duration(self, full_text: str) -> str:
        """Извлечение продолжительности обучения"""
        
        # Паттерны для поиска продолжительности
        duration_patterns = [
            r'срок\s+обучения[:\s]*(\d+)\s*(года?|лет)',
            r'продолжительность[:\s]*(\d+)\s*(года?|лет)',
            r'(\d+)\s*(года?|лет)\s*обучения',
            r'магистратура[:\s]*(\d+)\s*(года?|лет)'
        ]
        
        for pattern in duration_patterns:
            match = re.search(pattern, full_text.lower())
            if match:
                return f"{match.group(1)} года"
        
        return "2 года"  # Стандартная продолжительность магистратуры
    
    def _extract_admission_requirements(self, full_text: str, paragraphs: List[str]) -> List[str]:
        """Извлечение требований к поступлению"""
        
        requirements = []
        
        # Ключевые слова для поиска требований
        requirement_keywords = [
            'требования', 'поступление', 'вступительные', 'испытания',
            'экзамен', 'конкурс', 'отбор'
        ]
        
        text_lower = full_text.lower()
        
        # Ищем секции с требованиями
        for i, paragraph in enumerate(paragraphs):
            paragraph_lower = paragraph.lower()
            
            if any(keyword in paragraph_lower for keyword in requirement_keywords):
                # Собираем следующие несколько параграфов как требования
                for j in range(i, min(i + 3, len(paragraphs))):
                    next_paragraph = paragraphs[j].strip()
                    if len(next_paragraph) > 20 and not self._is_header(next_paragraph):
                        requirements.append(next_paragraph)
                break
        
        # Стандартные требования если ничего не найдено
        if not requirements:
            requirements = [
                "Высшее образование (диплом бакалавра или специалиста)",
                "Вступительные испытания по профилю программы",
                "Конкурсный отбор по результатам вступительных испытаний"
            ]
        
        return requirements
    
    def _extract_career_prospects(self, full_text: str, paragraphs: List[str]) -> List[str]:
        """Извлечение карьерных перспектив"""
        
        prospects = []
        
        # Ключевые слова для карьерных перспектив
        career_keywords = [
            'карьера', 'трудоустройство', 'работа', 'профессия', 
            'деятельность', 'специалист', 'должность'
        ]
        
        # Ищем секции с карьерными перспективами
        for i, paragraph in enumerate(paragraphs):
            paragraph_lower = paragraph.lower()
            
            if any(keyword in paragraph_lower for keyword in career_keywords):
                # Извлекаем должности и специальности
                for j in range(i, min(i + 3, len(paragraphs))):
                    next_paragraph = paragraphs[j]
                    career_items = self._extract_career_items(next_paragraph)
                    prospects.extend(career_items)
        
        # Стандартные перспективы на основе названия программы
        if not prospects:
            if "AI-продукт" in full_text:
                prospects = [
                    "Продуктовый менеджер AI-продуктов",
                    "AI Product Manager",
                    "Data Product Manager",
                    "ML Engineer",
                    "AI Consultant"
                ]
            else:
                prospects = [
                    "ML-инженер",
                    "Data Scientist",
                    "AI-разработчик",
                    "Исследователь в области ИИ",
                    "Аналитик данных"
                ]
        
        return list(set(prospects))  # Убираем дубликаты
    
    def _extract_career_items(self, text: str) -> List[str]:
        """Извлечение отдельных карьерных позиций из текста"""
        
        # Паттерны для поиска должностей
        job_patterns = [
            r'(\w+\s*){1,3}(?:инженер|менеджер|аналитик|разработчик|специалист|консультант)',
            r'(?:ML|AI|Data)\s+\w+',
            r'\w+\s+Scientist',
            r'Product\s+Manager'
        ]
        
        jobs = []
        for pattern in job_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                job = match.group().strip()
                if len(job) > 5:
                    jobs.append(job)
        
        return jobs
    
    def _extract_courses_from_document(self, full_text: str, paragraphs: List[str], 
                                     tables_data: List, program_name: str) -> List[Course]:
        """Извлечение курсов из документа"""
        
        courses = []
        program_id = "AI" if "интеллект" in program_name.lower() else "AI_Product"
        
        # Извлекаем курсы из таблиц
        for table in tables_data:
            table_courses = self._extract_courses_from_table(table, program_id)
            courses.extend(table_courses)
        
        # Если курсы не найдены в таблицах, ищем в тексте
        if not courses:
            courses = self._extract_courses_from_text(paragraphs, program_id)
        
        return courses
    
    def _extract_courses_from_table(self, table_data: List[List[str]], program_id: str) -> List[Course]:
        """Извлечение курсов из таблицы"""
        
        courses = []
        
        if not table_data or len(table_data) < 2:
            return courses
        
        # Предполагаем, что первая строка - заголовки
        headers = [cell.lower() for cell in table_data[0]]
        
        # Ищем индексы нужных колонок
        name_idx = self._find_column_index(headers, ['название', 'наименование', 'дисциплина'])
        credits_idx = self._find_column_index(headers, ['кредит', 'з.е', 'зачет'])
        semester_idx = self._find_column_index(headers, ['семестр', 'период'])
        
        # Обрабатываем строки с данными
        for row in table_data[1:]:
            if len(row) > max(name_idx or 0, credits_idx or 0, semester_idx or 0):
                
                name = row[name_idx] if name_idx is not None else ""
                credits = self._parse_credits(row[credits_idx] if credits_idx is not None else "3")
                semester = row[semester_idx] if semester_idx is not None else "1 семестр"
                
                # Фильтрация некачественных названий курсов
                if name and self._is_valid_course_name(name):
                    course = Course(
                        name=name.strip(),
                        description=f"Курс по программе {program_id}",
                        credits=credits,
                        semester=semester,
                        is_mandatory=True,  # По умолчанию считаем обязательным
                        program=program_id,
                        tags=self._generate_course_tags(name),
                        prerequisites=[]
                    )
                    courses.append(course)
        
        return courses
    
    def _is_valid_course_name(self, name: str) -> bool:
        """Проверка валидности названия курса"""
        name = name.strip()
        
        # Слишком короткие названия
        if len(name) < 5:
            return False
        
        # Слишком длинные названия (вероятно захватился лишний текст)
        if len(name) > 150:
            return False
        
        # Исключаем явно нерелевантные названия
        invalid_patterns = [
            'изначально я учился',
            'вступительные испытания',
            'основные даты',
            'показать все',
            'http',
            'www',
            'подробнее',
            'читать далее',
            '20241/5',  # Рейтинги и даты
            'никита борисов',  # Имена людей
            'ai talent hub'  # Названия программ
        ]
        
        name_lower = name.lower()
        for pattern in invalid_patterns:
            if pattern in name_lower:
                return False
        
        # Проверяем на наличие признаков настоящего курса
        course_indicators = [
            'анализ', 'программирование', 'алгоритм', 'метод', 'основы',
            'введение', 'технология', 'система', 'обработка', 'разработка',
            'математика', 'статистика', 'машинное', 'глубокое', 'обучение',
            'искусственный', 'интеллект', 'данные', 'python', 'практика',
            'теория', 'моделирование', 'проектирование', 'информатика'
        ]
        
        # Если есть хотя бы один индикатор курса - считаем валидным
        if any(indicator in name_lower for indicator in course_indicators):
            return True
        
        # Если название состоит только из заглавных букв - может быть аббревиатурой курса
        if name.isupper() and 10 <= len(name) <= 50:
            return True
        
        return False
    
    def _find_column_index(self, headers: List[str], keywords: List[str]) -> Optional[int]:
        """Поиск индекса колонки по ключевым словам"""
        for i, header in enumerate(headers):
            if any(keyword in header for keyword in keywords):
                return i
        return None
    
    def _parse_credits(self, credits_str: str) -> int:
        """Парсинг количества кредитов"""
        numbers = re.findall(r'\d+', credits_str)
        if numbers:
            credits = int(numbers[0])
            return credits if 1 <= credits <= 12 else 3
        return 3
    
    def _extract_courses_from_text(self, paragraphs: List[str], program_id: str) -> List[Course]:
        """Извлечение курсов из текста (если нет таблиц)"""
        
        courses = []
        
        # Ищем списки курсов в тексте
        for paragraph in paragraphs:
            # Ищем строки, которые могут быть названиями курсов
            if self._looks_like_course_name(paragraph):
                course = Course(
                    name=paragraph.strip(),
                    description=f"Курс по программе {program_id}",
                    credits=3,
                    semester="1 семестр",
                    is_mandatory=True,
                    program=program_id,
                    tags=self._generate_course_tags(paragraph),
                    prerequisites=[]
                )
                courses.append(course)
        
        return courses
    
    def _looks_like_course_name(self, text: str) -> bool:
        """Проверка, похож ли текст на название курса"""
        
        course_indicators = [
            'анализ', 'программирование', 'алгоритм', 'метод', 'основы',
            'введение', 'технология', 'система', 'обработка', 'разработка'
        ]
        
        text_lower = text.lower()
        
        return (20 <= len(text) <= 150 and
                any(indicator in text_lower for indicator in course_indicators) and
                not self._is_header(text))
    
    def _generate_course_tags(self, course_name: str) -> List[str]:
        """Генерация тегов для курса"""
        
        tags = []
        name_lower = course_name.lower()
        
        # Основные теги на основе ключевых слов
        tag_keywords = {
            'Machine Learning': [
                'машинное обучение', 'machine learning', 'ml', 'автоматическое машинное обучение',
                'применения машинного обучения', 'введение в мо', 'продвинутое мо'
            ],
            'Deep Learning': [
                'глубокое обучение', 'deep learning', 'нейронные сети', 'neural networks',
                'основы глубокого обучения', 'глубокие нейронные сети'
            ],
            'Computer Vision': [
                'компьютерное зрение', 'computer vision', 'cv', 'обработка изображений',
                'распознавание образов', 'image processing', 'генерация изображений',
                'обработка и генерация изображений', 'визуальный', 'изображение'
            ],
            'NLP': [
                'обработка естественного языка', 'natural language processing', 'nlp',
                'анализ текста', 'text mining', 'обработка языка', 'лингвистика',
                'языковые модели', 'текстовый анализ'
            ],
            'Data Science': [
                'data science', 'анализ данных', 'большие данные', 'big data',
                'статистика', 'analytics', 'математическая статистика',
                'анализ и обработка данных', 'визуализация данных'
            ],
            'Python': [
                'python', 'программирование на python', 'python backend',
                'разработка веб-приложений (python backend)', 'введение в мо (python)',
                'продвинутое мо (python)'
            ],
            'Programming': [
                'программирование', 'разработка', 'веб-приложений', 'backend',
                'программирование на с++', 'алгоритмы и структуры данных'
            ],
            'Research': [
                'исследования', 'research', 'научная работа', 'методология',
                'анализ', 'теория', 'фундаментальные'
            ],
            'Algorithms': [
                'алгоритм', 'структуры данных', 'алгоритмы и структуры данных',
                'оптимизация', 'complexity'
            ],
            'Math': [
                'математика', 'математическая', 'статистика', 'probability', 'вероятность',
                'линейная алгебра', 'linear algebra', 'математическая статистика'
            ],
            'Recommender Systems': [
                'рекомендательные системы', 'рекомендации', 'recommender systems',
                'персонализация', 'collaborative filtering'
            ],
            'Web Development': [
                'веб-разработка', 'web development', 'веб-приложения', 'backend',
                'frontend', 'разработка веб-приложений'
            ]
        }
        
        # Ищем совпадения с ключевыми словами
        for tag, keywords in tag_keywords.items():
            for keyword in keywords:
                if keyword in name_lower:
                    tags.append(tag)
                    break  # Достаточно одного совпадения для добавления тега
        
        # Убираем дубликаты и возвращаем
        return list(set(tags))
    
    def parse_all_docx_files(self, directory: str = ".") -> List[Program]:
        """Парсинг всех .docx файлов в директории"""
        
        programs = []
        
        # Ищем все .docx файлы
        docx_files = []
        for filename in os.listdir(directory):
            if filename.endswith('.docx') and not filename.startswith('~'):
                docx_files.append(os.path.join(directory, filename))
        
        logger.info(f"Найдено {len(docx_files)} документов для парсинга")
        
        # Парсим каждый файл
        for file_path in docx_files:
            program = self.parse_docx_file(file_path)
            if program:
                programs.append(program)
                logger.info(f"Успешно спарсена программа: {program.name}")
        
        return programs

# Пример использования
if __name__ == "__main__":
    parser = DocxParser()
    programs = parser.parse_all_docx_files(".")
    
    print(f"Найдено программ: {len(programs)}")
    for program in programs:
        print(f"\n=== {program.name} ===")
        print(f"Описание: {program.description[:100]}...")
        print(f"Продолжительность: {program.duration}")
        print(f"Курсов: {len(program.courses)}")
        print(f"Требования: {len(program.admission_requirements)}")
        print(f"Карьерные перспективы: {len(program.career_prospects)}") 